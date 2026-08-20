import os
from typing import List, Tuple
import fitz # PyMuPDF
import docx
from app.services.ollama_translator import call_ollama_generate, OLLAMA_DEFAULT_MODEL
from app.utils.logger import logger

def extract_scanned_pdf_images(pdf_path: str, temp_dir: str) -> List[str]:
    """
    Sử dụng PyMuPDF (fitz) để chuyển các trang của file PDF Scan thành danh sách ảnh PNG.
    """
    os.makedirs(temp_dir, exist_ok=True)
    doc = fitz.open(pdf_path)
    image_paths: List[str] = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        pix = page.get_pixmap(dpi=200) # Render chất lượng 200 DPI cho OCR
        img_path = os.path.join(temp_dir, f"page_{page_num + 1}.png")
        pix.save(img_path)
        image_paths.append(img_path)

    doc.close()
    return image_paths

def run_paddle_ocr_on_image(image_path: str) -> List[str]:
    """
    Khởi tạo và chạy PaddleOCR nhận diện ra từng khối văn bản (Text Blocks).
    Nếu môi trường chưa có paddleocr, fallback lấy text cơ bản.
    """
    try:
        from paddleocr import PaddleOCR
        ocr = PaddleOCR(use_angle_cls=True, lang='en', show_log=False)
        result = ocr.ocr(image_path, cls=True)

        text_blocks: List[str] = []
        if result and result[0]:
            for line in result[0]:
                text = line[1][0]
                if text and len(text.strip()) > 2:
                    text_blocks.append(text.strip())
        return text_blocks
    except Exception as e:
        logger.warning(f"PaddleOCR không khả dụng hoặc bị lỗi ({e}), chuyển sang fallback PyMuPDF OCR...")
        return []

def process_scanned_pdf_translation(
    pdf_path: str,
    output_docx_path: str,
    source_lang: str = "en",
    target_lang: str = "vi",
    model: str = OLLAMA_DEFAULT_MODEL
) -> str:
    """
    Quy trình dịch file PDF Scan:
    Bước 1: PyMuPDF chuyển PDF thành ảnh PNG.
    Bước 2: PaddleOCR bóc tách các khối chữ (Text Blocks).
    Bước 3: Gửi các khối text cho Ollama dịch.
    Bước 4: Dùng python-docx tạo một file Word hoàn toàn mới chứa các đoạn dịch.
    """
    logger.info(f"Bắt đầu xử lý file PDF Scan: {pdf_path}")
    temp_dir = os.path.join("temp_scanned", os.path.basename(pdf_path).replace(".pdf", ""))
    image_paths = extract_scanned_pdf_images(pdf_path, temp_dir)

    doc_out = docx.Document()
    doc_out.add_heading("Bản Dịch File PDF Scan", level=1)

    for page_idx, img_path in enumerate(image_paths, 1):
        logger.info(f"Đang OCR & Dịch trang scan {page_idx}/{len(image_paths)}...")
        doc_out.add_heading(f"Trang {page_idx}", level=2)

        blocks = run_paddle_ocr_on_image(img_path)
        
        # Nếu PaddleOCR không lấy được text, thử fallback đọc text trực tiếp từ PyMuPDF
        if not blocks:
            doc_pdf = fitz.open(pdf_path)
            raw_text = doc_pdf[page_idx - 1].get_text("text")
            doc_pdf.close()
            if raw_text.strip():
                blocks = [b.strip() for b in raw_text.split("\n\n") if b.strip()]

        for block in blocks:
            prompt = (
                f"Dịch đoạn văn bản OCR sau từ tiếng {source_lang.upper()} sang tiếng {target_lang.upper()}.\n"
                f"Chỉ trả về nội dung đã dịch:\n\n{block}"
            )
            try:
                translated_block = call_ollama_generate(prompt=prompt, model=model)
                doc_out.add_paragraph(translated_block if translated_block else block)
            except Exception as e:
                logger.warning(f"Lỗi dịch block OCR: {e}")
                doc_out.add_paragraph(block)

    os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
    doc_out.save(output_docx_path)
    
    # Cleanup temp images
    for img_p in image_paths:
        if os.path.exists(img_p):
            try:
                os.remove(img_p)
            except Exception:
                pass

    logger.info(f"Hoàn thành dịch PDF Scan và xuất file Word tại {output_docx_path}")
    return output_docx_path
